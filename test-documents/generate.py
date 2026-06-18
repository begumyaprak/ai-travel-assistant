from docx import Document

def create_doc(filename, title, content):
    doc = Document()
    doc.add_heading(title, 0)
    for section_title, paragraphs in content:
        doc.add_heading(section_title, 1)
        for para in paragraphs:
            doc.add_paragraph(para)
    doc.save(filename)
    print(f"Created: {filename}")

# 1. Barcelona Hotel Policy
create_doc("barcelona-hotel-policy.docx", "Barcelona Hotel Check-in & Policy Guide", [
    ("Check-in & Check-out", [
        "Standard check-in time is 3:00 PM. Early check-in before 3:00 PM is available upon request and subject to availability. An early check-in fee of 50 EUR may apply.",
        "Standard check-out time is 11:00 AM. Late check-out until 2:00 PM can be requested at the front desk for an additional fee of 30 EUR. Check-out after 2:00 PM will be charged as a full night.",
    ]),
    ("Cancellation Policy", [
        "Free cancellation is available up to 48 hours before the check-in date. Cancellations made within 48 hours of check-in will be charged one night's stay.",
        "No-show reservations will be charged the full booking amount.",
    ]),
    ("Amenities", [
        "The hotel offers complimentary Wi-Fi throughout the property. Breakfast is included in all room rates and is served from 7:00 AM to 10:30 AM in the main dining room.",
        "The rooftop pool is open from 9:00 AM to 9:00 PM daily. Towels are provided at the pool area. The hotel gym is open 24 hours and is accessible with your room key.",
        "Parking is available in the underground garage at a rate of 25 EUR per night.",
    ]),
    ("Pet Policy", [
        "Pets are welcome at the hotel. A non-refundable pet fee of 40 EUR per night applies. Maximum weight limit is 10 kg. Pets must be kept on a leash in all common areas.",
    ]),
    ("Payment & Deposits", [
        "A credit card is required at check-in for incidental charges. A security deposit of 200 EUR will be pre-authorized on your card upon arrival and released within 5 business days after check-out.",
        "Accepted payment methods: Visa, Mastercard, American Express, and cash in EUR.",
    ]),
])

# 2. Rome Destination Guide
create_doc("rome-destination-guide.docx", "Rome Travel Destination Guide", [
    ("Overview", [
        "Rome, the Eternal City, is the capital of Italy and one of the most historically rich destinations in the world. With over 2,500 years of history, Rome offers an unparalleled mix of ancient ruins, Renaissance art, and vibrant modern culture.",
        "The best time to visit Rome is from April to June and September to October when the weather is pleasant and crowds are smaller compared to peak summer months.",
    ]),
    ("Top Attractions", [
        "The Colosseum: Built in 70-80 AD, the Colosseum is the largest ancient amphitheater ever built. Tickets cost 16 EUR and should be booked in advance to avoid long queues. Opening hours are 9:00 AM to 7:00 PM.",
        "Vatican Museums & Sistine Chapel: Located within Vatican City, the museums house one of the world's greatest art collections. Tickets are 20 EUR and advance booking is strongly recommended. The museums are closed on Sundays except the last Sunday of each month when entry is free.",
        "Trevi Fountain: The largest Baroque fountain in Rome. Entry is free but the area is very crowded. Visit early morning (before 8:00 AM) for a quieter experience.",
        "Roman Forum: The center of ancient Roman public life. Combined tickets with the Colosseum are available. Open daily from 9:00 AM to sunset.",
    ]),
    ("Transportation", [
        "Rome's public transport includes metro, bus, and tram lines. A single ticket costs 1.50 EUR and is valid for 100 minutes. A 48-hour tourist pass costs 7 EUR and offers unlimited travel on all public transport.",
        "Taxis are metered. The minimum fare is 3.50 EUR during the day and 6.50 EUR at night. From Fiumicino Airport to the city center, the fixed fare is 48 EUR.",
        "Walking is the best way to explore the historic center. Most major attractions are within walking distance of each other.",
    ]),
    ("Food & Dining", [
        "Rome is famous for its pasta dishes including Cacio e Pepe, Carbonara, and Amatriciana. Avoid restaurants directly next to tourist attractions as they tend to be overpriced.",
        "A typical meal at a local trattoria costs 15-25 EUR per person including wine. Tipping is not mandatory but rounding up the bill is appreciated.",
        "Gelato shops (gelaterias) are everywhere. Authentic gelato is stored in metal containers with lids, not piled high in plastic bins.",
    ]),
    ("Practical Information", [
        "Currency: Euro (EUR). Most establishments accept credit cards but it is advisable to carry some cash for small purchases.",
        "Language: Italian. English is widely spoken in tourist areas and hotels.",
        "Emergency number: 112. Tourist police (Carabinieri) can be reached at 113.",
        "Dress code: When visiting churches and the Vatican, shoulders and knees must be covered. Scarves and shawls can be purchased near entrances.",
    ]),
])

# 3. Travel Insurance Policy
create_doc("travel-insurance-policy.docx", "TravelSafe Premium Insurance Policy", [
    ("Coverage Overview", [
        "TravelSafe Premium provides comprehensive travel insurance for international and domestic trips. This policy covers trip cancellation, medical emergencies, baggage loss, and travel delays.",
        "Maximum coverage amount: 50,000 USD per person per trip. Policy is valid for trips up to 90 days in duration.",
    ]),
    ("Medical Coverage", [
        "Emergency medical expenses are covered up to 50,000 USD. This includes hospitalization, surgery, and emergency dental treatment up to 500 USD.",
        "Medical evacuation and repatriation: covered up to 100,000 USD. In case of a medical emergency, contact our 24/7 assistance line at +1-800-TRAVEL-SAFE before seeking treatment when possible.",
        "Pre-existing medical conditions are not covered unless declared at the time of purchase and accepted in writing by TravelSafe.",
    ]),
    ("Trip Cancellation & Interruption", [
        "Trip cancellation coverage reimburses non-refundable trip costs up to 10,000 USD if you must cancel due to covered reasons including illness, injury, death of a family member, or natural disaster.",
        "Trip interruption coverage reimburses additional transportation costs to return home early plus unused non-refundable trip costs, up to 150% of the insured trip cost.",
        "Covered reasons for cancellation include: serious illness or injury of the insured or travel companion, death of an immediate family member, mandatory evacuation at destination, terrorist attack at destination.",
    ]),
    ("Baggage & Personal Belongings", [
        "Baggage loss or theft is covered up to 2,500 USD per person. Individual item limit is 500 USD. Electronics are covered up to 1,000 USD in total.",
        "Baggage delay: if your baggage is delayed more than 12 hours, you will be reimbursed up to 300 USD for essential purchases such as clothing and toiletries.",
        "To claim baggage loss, you must report the loss to the airline or relevant authority immediately and obtain a written report. Claims must be submitted within 90 days of the incident.",
    ]),
    ("Travel Delay", [
        "If your trip is delayed by more than 6 hours due to a covered reason (weather, mechanical breakdown, strike), you will be reimbursed up to 150 USD per day for meals and accommodation, maximum 600 USD.",
        "Flight cancellation by the airline is covered. You must request compensation from the airline first before claiming from TravelSafe.",
    ]),
    ("Claims Process", [
        "To file a claim, visit our website at www.travelsafe.com/claims or call +1-800-TRAVEL-SAFE. All claims must be submitted within 90 days of the incident.",
        "Required documentation: completed claim form, proof of payment, receipts for expenses claimed, police report (if applicable), medical records (for medical claims), airline documentation (for delay/cancellation claims).",
        "Claims are typically processed within 10-15 business days of receiving all required documentation.",
    ]),
])

# 4. Tokyo Destination Guide
create_doc("tokyo-destination-guide.docx", "Tokyo Travel Destination Guide", [
    ("Overview", [
        "Tokyo, the capital of Japan, is one of the world's most dynamic and fascinating cities. A seamless blend of ultramodern technology and ancient tradition, Tokyo offers an extraordinary travel experience unlike anywhere else on earth.",
        "The best time to visit Tokyo is during spring (late March to early May) for cherry blossoms, or autumn (October to November) for colorful foliage. Summers are hot and humid; winters are mild but can be cold.",
    ]),
    ("Top Attractions", [
        "Senso-ji Temple (Asakusa): Tokyo's oldest and most iconic temple, founded in 645 AD. Entry is free and the temple is open 24 hours. The Nakamise shopping street leading to the temple is perfect for souvenirs.",
        "Shibuya Crossing: One of the busiest pedestrian crossings in the world. Best viewed from the Starbucks or Mag's Park observation deck above the intersection. Completely free to experience.",
        "Shinjuku Gyoen National Garden: A beautiful garden combining Japanese, French, and English landscape styles. Entrance fee is 500 JPY. Open Tuesday to Sunday, 9:00 AM to 4:30 PM. Especially stunning during cherry blossom season.",
        "teamLab Planets (Toyosu): An immersive digital art museum. Tickets cost 3,200 JPY and must be booked online in advance. Open daily 9:00 AM to 9:00 PM.",
        "Tsukiji Outer Market: The world-famous fish market. Best visited early morning (around 6:00 AM) for the freshest sushi breakfast. Free to enter; budget 2,000–4,000 JPY for a meal.",
    ]),
    ("Transportation", [
        "Tokyo has one of the world's most efficient public transport systems. The IC card (Suica or Pasmo) can be used on all trains, subways, and buses. Load it at any station kiosk.",
        "A single metro ride costs between 170 and 320 JPY depending on distance. Day passes are available for 600 JPY (Tokyo Metro only) or 1,000 JPY (all lines).",
        "From Narita Airport to central Tokyo: Narita Express (N'EX) takes 60 minutes and costs 3,070 JPY. Limousine bus takes 90 minutes and costs 3,200 JPY. Taxis are available but expensive — expect 20,000–25,000 JPY.",
        "From Haneda Airport to central Tokyo: Tokyo Monorail or Keikyu Line takes 20–30 minutes and costs 500–700 JPY.",
    ]),
    ("Food & Dining", [
        "Tokyo has more Michelin-starred restaurants than any other city in the world. However, excellent food is available at every price point.",
        "Ramen, sushi, tempura, yakitori, and tonkatsu are must-try dishes. Conveyor belt sushi (kaiten-zushi) is a fun and affordable option at 150–500 JPY per plate.",
        "Convenience stores (7-Eleven, FamilyMart, Lawson) offer surprisingly high-quality meals for 500–800 JPY. Onigiri, sandwiches, and hot foods are available 24 hours.",
        "Tipping is not customary in Japan and can even be considered rude. Service charges are never added to bills.",
    ]),
    ("Practical Information", [
        "Currency: Japanese Yen (JPY). Japan is still largely cash-based — always carry sufficient yen. 7-Eleven ATMs reliably accept foreign cards.",
        "Language: Japanese. English signage is available at major train stations and tourist areas. Translation apps (Google Translate camera mode) are very helpful.",
        "Wi-Fi: Pocket Wi-Fi devices can be rented at the airport (around 700 JPY/day). Free Wi-Fi is available at most convenience stores and major tourist spots.",
        "Emergency number: 110 (police), 119 (ambulance/fire). The Japan Visitor Hotline (+81-50-3816-2787) provides 24/7 assistance in English.",
        "Etiquette: Remove shoes when entering traditional restaurants or ryokan. Avoid eating or drinking while walking. Queuing is strictly observed.",
    ]),
])

print("\nAll documents created successfully in test-documents/")